import base64, io, time
import fitz                      # pip install PyMuPDF
from openai import OpenAI
from docx import Document        # <-- 1. add this  (pip install python-docx)

client = OpenAI(api_key="EMPTY", base_url="http://163.128.34.25:5000/v1", timeout=3600)
doc = fitz.open("swimming.pdf")
pages = []

for i, page in enumerate(doc, 1):
    pix = page.get_pixmap(matrix=fitz.Matrix(200/72, 200/72), alpha=False)
    b64 = base64.b64encode(pix.tobytes("png")).decode()
    messages = [{
        "role": "user",
        "content": [
            {"type": "text", "text": "<image>\nConvert the document to markdown.","skip_special_tokens": True,},
            {"type": "image_url",
             "image_url": {"url": f"data:image/png;base64,{b64}"}},
        ],
    }]
    start = time.time()
    response = client.chat.completions.create(
        model="baidu/Unlimited-OCR",
        messages=messages,
        max_tokens=8192,
        temperature=0.0,
        extra_body={
            "skip_special_tokens": False,
            "vllm_xargs": {"ngram_size": 35, "window_size": 128},
        },
    )
    text = response.choices[0].message.content
    print(f"page {i}: {time.time() - start:.2f}s, {len(text)} chars")
    pages.append(text)

doc.close()

# ---- 2. turn the model's markdown into Word paragraphs --------------------
def add_line(word_doc, line):
    t = line.rstrip()
    if not t:
        return
    if t.startswith("### "):
        word_doc.add_heading(t[4:], level=3)
    elif t.startswith("## "):
        word_doc.add_heading(t[3:], level=2)
    elif t.startswith("# "):
        word_doc.add_heading(t[2:], level=1)
    elif t.lstrip().startswith(("- ", "* ")):
        word_doc.add_paragraph(t.lstrip()[2:], style="List Bullet")
    else:
        word_doc.add_paragraph(t)

# ---- 3. write the .docx instead of the .md -------------------------------
word_doc = Document()
for i, text in enumerate(pages, 1):
    for line in text.split("\n"):
        add_line(word_doc, line)
    if i < len(pages):
        word_doc.add_page_break()
word_doc.save("sample.docx")
print("saved: sample.docx")
